import requests  
from bs4 import BeautifulSoup  
import os  
import urllib.request 
from docx import Document  

def get_dirname(url):
    # 发送HTTP请求获取网页内容  
    response = requests.get(url)
    response.encoding = 'utf-8'  
    # 使用BeautifulSoup解析网页内容  
    soup = BeautifulSoup(response.text, 'html.parser')  
    # 获取标题  
    title = soup.find('h2', class_='title_con').text
    
    return title

def save_content_doc(url, output_dir):
    if not os.path.exists(output_dir):  # 如果输出目录不存在，创建它  
        os.makedirs(output_dir)
    # 发送HTTP请求获取网页内容  
    response = requests.get(url)
    response.encoding = 'utf-8'  
    # 使用BeautifulSoup解析网页内容  
    soup = BeautifulSoup(response.text, 'html.parser')  
    # 获取标题  
    title = soup.find('h2', class_='title_con').text  
    # 获取正文内容，这里以第一个div元素为例  
    content = soup.find('div', class_='my_doccontent').text  
    # 创建Word文档  
    doc = Document()  
    # 添加标题和正文到文档  
    doc.add_heading(title, level=1)  
    doc.add_paragraph(content)  
  
    filename = os.path.join(output_dir, f"{title}.docx")  # 输出文件名，包含完整路径
    # 保存Word文档  
    doc.save(filename)

def extract_filename_from_html(url):
    response = requests.get(url)
    response.encoding = 'utf-8'  
    soup = BeautifulSoup(response.text, 'html.parser')  
    appendix1_element = soup.find('span', {'id': 'appendix1'})
    
    file_names = []

    if appendix1_element:  
        for a_element in appendix1_element.find_all('a'):
            name = a_element.text.strip()
            file_names.append(name)  

        print(file_names)
        return file_names  
    return None 
  
def get_pdf_links(url):  
    response = requests.get(url)  
    soup = BeautifulSoup(response.text, 'html.parser')  
  
    pdf_links = []  
  
    for link in soup.find_all('a'):  # 查找所有超链接  
        href = link.get('href')  # 获取链接  
        if href and 'pdf' in href.lower():  # 如果链接是PDF文件  
            pdf_links.append(href)  # 添加到列表中  

    return pdf_links  
  
def download_pdf(pdf_links, output_dir, url_prefix, filenames):
    if not os.path.exists(output_dir):  # 如果输出目录不存在，创建它  
        os.makedirs(output_dir)

    for link, filename in zip(pdf_links, filenames):
        url_pdf = link.lstrip('.')
        url = url_prefix + url_pdf  # 请替换为实际的链接
        # url_pdf = url_pdf.lstrip('/')  
        file_name = os.path.join(output_dir, filename)  # 输出文件名，包含完整路径
        print(file_name)  
        with urllib.request.urlopen(url) as url:  
            with open(file_name, 'wb') as file:  
                file.write(url.read())  # 下载文件并保存到指定路径  

#<li><a href="./202311/t20231124_3918214.htm" target="_blank" title="关于印发《代理记账基础工作规范（试行）》的通知">关于印发《代理记账基础工作规范（试行）》的通知</a><span>2023-11-24</span></li>
def get_urls(url):
    response = requests.get(url)
    response.encoding = 'utf-8'
    soup = BeautifulSoup(response.text, 'html.parser')
    # ul class="liBox"
    liBox_element = soup.find('ul', {'class': 'liBox'})
    # print(liBox_element)
    urls = []
    if liBox_element:
        for url in liBox_element.find_all('a'):  # 查找所有超链接  
            href = url.get('href')  # 获取链接  
            if href and 'htm' in href.lower():  # 如果链接是htm文件  
                urls.append(href)  # 添加到列表中  
        # print(urls)       
        return urls
    
    return None

def get_pagerji(url):
    response = requests.get(url)

    soup = BeautifulSoup(response.text, 'html.parser')
    # ul class="liBox"
    urls = []
    for pagerji_element in soup.find_all('p', {'class': 'pagerji'}):
        print(pagerji_element)
    
        if pagerji_element:
            for url in pagerji_element.find_all('a'):
                href = url.get('href')
                if href and 'htm' in href.lower():
                    urls.append(href)
            print(urls)        
    return urls
    
    return None
    
# 测试代码  

# output_dir = "pdfs"  # 请替换为实际的下载文件保存路径
# 定义处理urls的函数  
def process_urls(urls): 
# url = "http://kjs.mof.gov.cn/zhengcefabu/202311/t20231124_3918214.htm"  # 请替换为实际的网页链接
    for url in urls:
        output_dir = get_dirname(url)  
        save_content_doc(url, output_dir)
        filenames = extract_filename_from_html(url)  

        last_slash_index = url.rfind("/")  
        field_before_last_slash = url[:last_slash_index]   

        pdf_links = get_pdf_links(url)  
        download_pdf(pdf_links, output_dir, field_before_last_slash, filenames)
        
def main():
    # pagerji_urls = get_pagerji("http://kjs.mof.gov.cn/zhengcefabu/")
    # print(pagerji_urls)
    
    urls = get_urls("http://kjs.mof.gov.cn/zhengcefabu/index.htm")
    
    final_urls = []
    
    for url in urls:
        sub_url = url.lstrip('./')
        url = "http://kjs.mof.gov.cn/zhengcefabu/" + sub_url
        final_urls.append(url)

    process_urls(final_urls)

if __name__ == "__main__":  
    main()